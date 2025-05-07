import os
from docx import Document
from docx.shared import Pt

def find_and_process_table(doc_path):
    if not os.path.exists(doc_path):
        print(f"❌ 错误：文件 '{doc_path}' 不存在！")
        return

    document = Document(doc_path)

    # 查找标题“制造策略速查表”
    target_paragraph_index = None
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.startswith('Heading') and "制造策略速查表" in paragraph.text:
            target_paragraph_index = i
            print(f"✅ 找到标题 '{paragraph.text}' 在第 {i} 段")
            break
    if target_paragraph_index is None:
        print("未找到标题 '制造策略速查表'")
        return

    # 从该段落后开始查找第一个表格
    table_found = None
    for element in document.element.body[target_paragraph_index:]:
        if element.tag.endswith('tbl'):
            print("✅ 找到表格元素")
            table_found = element
            break

    if table_found is None:
        print("未找到标题下的表格")
        return

    # 将表格对象转换为 Table 类型
    table = None
    for t in document.tables:
        if t._element == table_found:
            table = t
            break

    if table is None:
        print("无法将XML表格转换为Table对象")
        return

    # 获取表头并打印表格内容
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    types = headers[2:]  # A1, A2, B1, B2, C1, C2, D1, D2

    print("表格内容如下：")
    for i, row in enumerate(table.rows):
        print(f"Row {i}: {[cell.text.strip() for cell in row.cells]}")

    # 构造嵌套结构：{type_name: {process: {sub_process: value}}}
    data_by_type = {t: {} for t in types}

    current_process = ""
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]

        # 如果第一列有值，则更新当前主工序
        if cells[0]:
            current_process = cells[0]

        sub_process = cells[1] if len(cells) > 1 else ""

        # 遍历每个类型列
        for col_idx, detail in enumerate(cells[2:], start=2):
            type_name = headers[col_idx]

            if detail and current_process and sub_process:
                if current_process not in data_by_type[type_name]:
                    data_by_type[type_name][current_process] = {}

                data_by_type[type_name][current_process][sub_process] = detail

    # 生成描述性文字并插入到文档中
    for type_name in types:
        description_text = f"{type_name}：\n"
        for process, sub_processes in data_by_type[type_name].items():
            description_text += f"- {process}：\n"
            for sub, val in sub_processes.items():
                description_text += f"  - {sub}：{val}\n"

        p = document.add_paragraph()
        run = p.add_run(description_text)
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(10)

    # 构造新文件路径并保存
    new_doc_path = os.path.join(
        os.path.dirname(doc_path),
        "modified_" + os.path.basename(doc_path)
    )

    document.save(new_doc_path)
    print(f"✅ 已成功生成新文档：{new_doc_path}")

# 示例调用
doc_path = "/Users/oliver/Desktop/云尖信息poc/网络/网络产品制造策略技术部分.docx"
find_and_process_table(doc_path)