from docx import Document
from docx.oxml.ns import qn
import re

def is_heading(paragraph):
    """判断段落是否为标题（如 Heading 1, Heading 2 等）"""
    return paragraph.style.name.startswith('Heading')

def find_next_non_heading_text(paragraphs, index):
    """
    在指定索引之后查找第一个非空且非标题的段落文本。
    如果遇到的第一个非空段落是标题，则认为图片没有名字。
    """
    for i in range(index + 1, len(paragraphs)):
        para_text = paragraphs[i].text.strip()
        if para_text:
            if is_heading(paragraphs[i]):
                print("提示: 第一个非空段落是标题，认为该图片没有名字")
                return "未找到相关文字"
            else:
                print(f"提示: 找到了非标题的后续文字：{para_text}")
                return para_text
    
    print("提示: 未找到非标题的后续文字")
    return "未找到相关文字"

def extract_images(doc_path):
    document = Document(doc_path)
    images_info = []
    img_counter = 0

    for i, paragraph in enumerate(document.paragraphs):
        # 检查当前段落是否包含图片
        if paragraph._element.findall('.//' + qn('w:drawing')):
            img_counter += 1
            
            # 获取图片下方的第一个非标题非空段落作为名字
            context_text = find_next_non_heading_text(document.paragraphs, i)

            # 使用找到的文本作为图片的名字（过滤非法字符）
            img_name = re.sub(r'[\\/:*?"<>|]', '_', context_text)[:50] 
            images_info.append({"name": img_name, "context": context_text})

    print(f"🖼️ 文档中共找到 {img_counter}张图片。")
    for img_info in images_info:
        print(f"🎯 成功定位图片，图片名字是：{img_info['name']}，附近的文字为：{img_info['context']}")

# 使用函数，指定你的.docx文档路径
extract_images('file/0507【输出-生产测试方案-案例-推荐】产品生产测试方案.docx')