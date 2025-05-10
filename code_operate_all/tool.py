from docx import Document
from docx.oxml.ns import qn
import re
import os


def is_heading(paragraph):
    return paragraph.style.name.startswith('Heading')


def find_next_non_heading_text(paragraphs, index):
    for i in range(index + 1, len(paragraphs)):
        para_text = paragraphs[i].text.strip()
        if para_text:
            if is_heading(paragraphs[i]):
                print("提示: 第一个非空段落是标题，认为该图片没有名字")
                return "未找到相关文字"
            else:
                print(f"提示: 找到了非标题的后续文字：{para_text}")
                return para_text
    print("提示: 未找到非标题的后续文字")
    return "未找到相关文字"


def extract_images_and_find_target(doc_path, target_title):
    document = Document(doc_path)
    img_counter = 0
    found_img_data = None
    images_info = []

    for i, paragraph in enumerate(document.paragraphs):
        # 检查当前段落是否包含图片
        drawings = paragraph._element.findall('.//' + qn('w:drawing'))
        if not drawings:
            continue

        # 获取图片下方的第一个非标题非空段落作为名字
        context_text = find_next_non_heading_text(document.paragraphs, i)

        if context_text == "未找到相关文字":
            print("⚠️ 跳过无有效说明文字的图片")
            continue

        img_counter += 1
        img_name = re.sub(r'[\\/:*?"<>|]', '_', context_text)[:50] + f"_img_{img_counter}"
        img_info = {"name": img_name, "context": context_text}
        images_info.append(img_info)

        # 判断是否是我们要找的目标图片
        if target_title.strip() in context_text.strip():
            print(f"✅ 找到了目标图片：“{target_title}”，正在尝试提取并分析...")

            image_part = None
            for drawing in drawings:
                # 使用 qn() 替代命名空间查找
                blips = drawing.findall('.//' + qn('a:blip'))  # a 是命名空间前缀，blip 是标签名
                if blips:
                    rId = blips[0].get(qn('r:embed'))  # 获取图片关系ID
                    if rId and rId in paragraph.part.rels:
                        image_part = paragraph.part.rels[rId].target_part
                        break

            if image_part:
                found_img_data = image_part.image.blob
                print("🖼️ 图片已读取到内存中")

                # 保存图片到当前文件夹
                img_format = image_part.image.ext
                img_filename = f"{img_info['name']}.{img_format}"
                with open(img_filename, "wb") as img_file:
                    img_file.write(found_img_data)
                print(f"💾 图片已保存为 {img_filename}")
            else:
                print("❌ 无法提取图片数据。")

    print(f"🖼️ 文档中共找到 {img_counter} 张图片。")
    for img_info in images_info:
        print(f"🎯 成功定位图片，图片名字是：{img_info['name']}，附近的文字为：{img_info['context']}")

    if not found_img_data:
        print("⚠️ 未在文档中找到目标图片。")

    print("found_img_data已经找到")
    return found_img_data